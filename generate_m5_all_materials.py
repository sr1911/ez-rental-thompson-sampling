"""Generate M5_All_Materials.docx - All English version."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15

def add_heading_styled(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
    return h

# ============ TITLE ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("M5 Project Paper - Complete Reference Materials")
run.font.size = Pt(16)
run.font.bold = True
run.font.name = 'Times New Roman'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("ISBA 2415, Reinforcement Learning")
run.font.size = Pt(11)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

# ============ PART 1 ============
add_heading_styled("Part 1: Assignment Instructions")
info = [
    ("Course", "ISBA 2415, Reinforcement Learning"),
    ("Assignment", "M5: Project Paper (Group)"),
    ("Due", "Mar 15 by 11:59pm"),
    ("Points", "300"),
    ("Submission", "File upload (only one member needs to submit)"),
    ("Plagiarism Check", "Unicheck enabled"),
]
for label, val in info:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run = p.add_run(val)
    run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Tasks")
run.font.bold = True
run.font.name = 'Times New Roman'
p = doc.add_paragraph()
p.add_run(
    "Implement the exploration versus exploitation strategy that you proposed in your project proposal "
    "for the EZ Car Rental case and answer the following questions:"
).font.name = 'Times New Roman'
p = doc.add_paragraph()
p.add_run("1. Were you able to learn the probability distribution that you were simulating your "
          "environment from?").font.name = 'Times New Roman'
p = doc.add_paragraph()
p.add_run("2. Were you able to uncover a pricing policy given the simulated environment?").font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Requirements")
run.font.bold = True
run.font.name = 'Times New Roman'
reqs = [
    "The group submissions should be approximately two pages in length.",
    "Please upload your code/notebooks on GitHub and provide a link to it.",
    "The paper should be supported by the use of charts and tables.",
    "Where applicable, please use a consistent referencing system such as APA, the Harvard system, the IEEE system, etc.",
    "Only one member from each group is required to submit this assignment.",
    "Please note that Unicheck has been enabled for this assignment.",
]
for i, r in enumerate(reqs):
    p = doc.add_paragraph()
    p.add_run(f"{i+1}. {r}").font.name = 'Times New Roman'

# ============ PART 2 ============
add_heading_styled("Part 2: M2 Proposal Recap")
p = doc.add_paragraph()
p.add_run("Below is a summary of what we proposed in our M2 submission:").font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Problem")
run.font.bold = True
run.font.name = 'Times New Roman'
p = doc.add_paragraph()
p.add_run(
    "EZ Car Rental operates a contactless car rental service across multiple U.S. cities and needs "
    "to learn optimal pricing in new markets with no historical pricing data. Supply is uncontrollable; "
    "we can only influence demand through pricing."
).font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run("DP Formulation")
run.font.bold = True
run.font.name = 'Times New Roman'
formulas = [
    "State Space: s = (city, time_category, utilization_level), 5 x 2 x 2 = 20 states",
    "Action Space: A = {$3, $5, $7, $9, $11}/hr, 5 prices",
    "Reward: Y ~ Bernoulli(theta), E[R(s,p)] = p * theta_{s,p} * A(s)",
    "Bellman Equation: V*(s) = max_p { E[R(s,p)] + gamma * sum P(s'|s,p) * V*(s') }",
]
for fm in formulas:
    p = doc.add_paragraph()
    p.add_run(fm).font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run("Exploration Strategy")
run.font.bold = True
run.font.name = 'Times New Roman'
p = doc.add_paragraph()
p.add_run(
    "Thompson Sampling with Beta-Bernoulli conjugate model. Each state-action pair maintains "
    "Beta(alpha, beta) posterior over the unknown booking probability."
).font.name = 'Times New Roman'

# ============ PART 3 ============
add_heading_styled("Part 3: Implementation Summary")
p = doc.add_paragraph()
run = p.add_run("What We Built")
run.font.bold = True
run.font.name = 'Times New Roman'
items = [
    "Simulator: Computed true booking probabilities from the Journeys CSV for each state-action pair. The agent has no direct access to this data.",
    "Thompson Sampling Agent: Beta-Bernoulli model with a Beta posterior for each of the 100 state-action pairs.",
    "Epsilon-Greedy Agent: epsilon = 0.15 as a comparison baseline.",
    "Oracle: Assumes full knowledge of true probabilities and computes the optimal policy as a benchmark.",
    "Ran 5,000 episodes with the same random seed to ensure a fair comparison.",
]
for i, item in enumerate(items):
    p = doc.add_paragraph()
    p.add_run(f"{i+1}. {item}").font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Key Results")
run.font.bold = True
run.font.name = 'Times New Roman'

table = doc.add_table(rows=6, cols=3, style='Light Shading Accent 1')
headers = ["Metric", "Thompson Sampling", "Epsilon-Greedy"]
data = [
    ["Total Revenue", "$19,519", "$19,550"],
    ["Revenue Ratio vs Oracle", "90.8%", "90.0%"],
    ["Total Regret", "$1,989", "$2,170"],
    ["Policy Accuracy", "20/20 (100%)", "19/20 (95%)"],
    ["Mean Abs Error", "0.108", "N/A"],
]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        cell = table.rows[r+1].cells[c]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Answers to Assignment Questions")
run.font.bold = True
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
run = p.add_run("Q1: Were we able to learn the probability distribution? ")
run.font.bold = True
run.font.name = 'Times New Roman'
p.add_run(
    "Yes. The mean absolute error between learned and true booking probabilities was 0.108 "
    "across all 100 state-action pairs. The probabilities that matter most for pricing decisions "
    "were estimated accurately."
).font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Q2: Were we able to uncover a pricing policy? ")
run.font.bold = True
run.font.name = 'Times New Roman'
p.add_run(
    "Yes. Thompson Sampling matched the oracle's optimal price in all 20 states (100% accuracy). "
    "Epsilon-greedy matched 19 out of 20 (95%). The learned policy is intuitive: lower prices for "
    "off-peak/low-utilization states and higher prices for peak/high-utilization states."
).font.name = 'Times New Roman'

# ============ PART 4 ============
add_heading_styled("Part 4: Charts")
charts = [
    ("Figure 1: Cumulative Regret - Thompson Sampling vs Epsilon-Greedy", "chart_regret.png"),
    ("Figure 2: Policy Accuracy Over Training", "chart_policy_accuracy.png"),
    ("Figure 3: Revenue Comparison Across Methods", "chart_revenue.png"),
    ("Figure 4: Pricing Policy Comparison Across All States", "chart_policy.png"),
]
for caption, filename in charts:
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    chart_path = os.path.join(BASE, filename)
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p2 = doc.add_paragraph()
        p2.add_run(f"[Chart not found: {filename}]").font.name = 'Times New Roman'
    doc.add_paragraph()

# ============ PART 5 ============
add_heading_styled("Part 5: Submission Checklist")

p = doc.add_paragraph()
run = p.add_run("Upload to Canvas:")
run.font.bold = True
run.font.name = 'Times New Roman'
p = doc.add_paragraph()
p.add_run("M5_Project_Paper.docx (make sure to fill in the GitHub link first)").font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Upload to GitHub:")
run.font.bold = True
run.font.name = 'Times New Roman'
gh_files = [
    "EZ_Car_Rental_Thompson_Sampling.ipynb",
    "data/journeys.csv",
]
for gf in gh_files:
    p = doc.add_paragraph()
    p.add_run(gf).font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Submission Steps:")
run.font.bold = True
run.font.name = 'Times New Roman'
steps = [
    "Create a new GitHub repo (e.g., ez-car-rental-rl)",
    "Upload the notebook + data/journeys.csv",
    "Open M5_Project_Paper.docx and replace [Insert GitHub link here] with your repo link",
    "Submit M5_Project_Paper.docx on Canvas",
]
for i, step in enumerate(steps):
    p = doc.add_paragraph()
    p.add_run(f"Step {i+1}: {step}").font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("File List:")
run.font.bold = True
run.font.name = 'Times New Roman'
files = [
    "M5_Project_Paper.docx - The paper to submit on Canvas",
    "EZ_Car_Rental_Thompson_Sampling.ipynb - Jupyter notebook for GitHub",
    "data/journeys.csv - Source data for the notebook",
    "ez_rental_thompson_sampling.py - Standalone Python script (backup)",
    "generate_charts.py - Chart generation script (backup)",
    "generate_m5_paper.py - Paper generation script (backup)",
    "M5_All_Materials.docx - This reference document (do not submit)",
]
for fl in files:
    p = doc.add_paragraph()
    p.add_run(fl).font.name = 'Times New Roman'

output_path = os.path.join(BASE, "M5_All_Materials.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
